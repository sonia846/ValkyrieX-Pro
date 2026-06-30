#!/usr/bin/env python3
"""
ValkyrieX Pro — Auto Report Generator
Author: Sonia
Description: Scan results ko TXT, HTML, JSON, DOCX format mein 
             professional reports mein convert karta hai.
"""

import json
import os
from datetime import datetime

AUTHOR_SIGNATURE = b"\x53\x6F\x6E\x69\x61"  # Sonia

def _verify_author():
    if AUTHOR_SIGNATURE.decode('utf-8') != "Sonia":
        raise PermissionError("ValkyrieX: Unauthorized execution blocked.")


def generate_all_reports(scan_results: dict, ai_analysis: dict, target: str, base_dir: str = "reports"):
    """Sab formats mein reports generate karo."""
    _verify_author()
    
    timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M-%S")
    safe_target = target.replace("https://", "").replace("http://", "").replace("/", "_")
    base_name = f"sonia_{safe_target}_{timestamp}"
    
    generated = []
    
    json_path = os.path.join(base_dir, "json", f"{base_name}.json")
    _write_json(scan_results, ai_analysis, json_path)
    generated.append(json_path)
    
    txt_path = os.path.join(base_dir, "txt", f"{base_name}.txt")
    _write_txt(scan_results, ai_analysis, target, txt_path)
    generated.append(txt_path)
    
    html_path = os.path.join(base_dir, "output", f"{base_name}.html")
    _write_html(scan_results, ai_analysis, target, html_path)
    generated.append(html_path)
    
    docx_path = os.path.join(base_dir, "docx", f"{base_name}.docx")
    _write_docx(scan_results, ai_analysis, target, docx_path)
    generated.append(docx_path)
    
    return generated


def _write_json(scan_results, ai_analysis, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    combined = {
        "author": "Sonia",
        "tool": "ValkyrieX Pro",
        "scan_results": scan_results,
        "ai_analysis": ai_analysis
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(combined, f, indent=4, default=str)


def _write_txt(scan_results, ai_analysis, target, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    
    vulns = [m for m in scan_results.get("module_outputs", []) 
             if "Vulnerable" in m.get("status", "") or "Alert" in m.get("status", "")]
    
    with open(path, 'w', encoding='utf-8') as f:
        f.write("=" * 70 + "\n")
        f.write("        VALKYRIEX PRO — SECURITY AUDIT REPORT\n")
        f.write("        Author: Sonia | Tool: ValkyrieX Pro\n")
        f.write("=" * 70 + "\n\n")
        f.write(f"[*] TARGET     : {target}\n")
        f.write(f"[*] SCAN TIME  : {scan_results.get('scan_time', 'N/A')} UTC\n")
        f.write(f"[*] TOTAL VULNS: {len(vulns)}\n")
        f.write(f"[*] REPORT BY  : Sonia\n")
        f.write("\n" + "=" * 70 + "\n")
        f.write("VULNERABILITY FINDINGS\n")
        f.write("=" * 70 + "\n\n")
        
        if not vulns:
            f.write("  [✓] No vulnerabilities detected in this scan.\n")
        else:
            for i, v in enumerate(vulns, 1):
                f.write(f"[{i:02d}] MODULE  : {v['module'].upper()}\n")
                f.write(f"     STATUS  : {v['status']}\n")
                f.write(f"     SUMMARY : {v['summary']}\n")
                f.write("-" * 70 + "\n\n")
        
        if ai_analysis and "ai_analysis" in ai_analysis:
            ai = ai_analysis["ai_analysis"]
            f.write("\n" + "=" * 70 + "\n")
            f.write("AI ANALYSIS SUMMARY\n")
            f.write("=" * 70 + "\n\n")
            f.write(f"  {ai.get('executive_summary', 'N/A')}\n")
            f.write(f"\n  Estimated Bounty: {ai.get('estimated_bounty_range', 'N/A')}\n")
        
        f.write("\n" + "=" * 70 + "\n")
        f.write("  END OF REPORT — ValkyrieX Pro by Sonia\n")
        f.write("=" * 70 + "\n")


def _write_html(scan_results, ai_analysis, target, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    
    vulns = [m for m in scan_results.get("module_outputs", []) 
             if "Vulnerable" in m.get("status", "") or "Alert" in m.get("status", "")]
    safe = [m for m in scan_results.get("module_outputs", []) 
            if m.get("status") == "Safe / Clean"]
    
    vuln_cards = ""
    for v in vulns:
        vuln_cards += f"""
        <div class="vuln-card">
            <div class="vuln-header">
                <span class="badge badge-danger">VULNERABLE</span>
                <strong>{v['module'].upper().replace('_', ' ')}</strong>
            </div>
            <div class="vuln-body">
                <p>{v.get('summary', 'N/A')}</p>
                <small>🕐 {v.get('timestamp', 'N/A')}</small>
            </div>
        </div>"""
    
    if not vuln_cards:
        vuln_cards = '<div class="no-vulns">✅ No vulnerabilities detected.</div>'
    
    ai_summary = ""
    if ai_analysis and "ai_analysis" in ai_analysis:
        ai = ai_analysis.get("ai_analysis", {})
        ai_summary = f"""
        <div class="ai-box">
            <h3>🧠 AI Analysis</h3>
            <p>{ai.get('executive_summary', 'No AI summary available.')}</p>
            <p><strong>💰 Estimated Bounty:</strong> {ai.get('estimated_bounty_range', 'N/A')}</p>
        </div>"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ValkyrieX Pro Report — {target}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ background: #0a0a0f; color: #e0e0e0; font-family: 'Courier New', monospace; }}
        .header {{ background: linear-gradient(135deg, #1a0030, #0d001a); padding: 40px; text-align: center; border-bottom: 2px solid #7c3aed; }}
        .header h1 {{ font-size: 2.5em; color: #a855f7; text-shadow: 0 0 20px #7c3aed; }}
        .header .author {{ color: #f0abfc; margin-top: 10px; font-size: 1.1em; }}
        .header .target {{ color: #86efac; margin-top: 8px; }}
        .stats {{ display: flex; gap: 20px; padding: 30px; justify-content: center; flex-wrap: wrap; }}
        .stat-card {{ background: #12001f; border: 1px solid #7c3aed; border-radius: 10px; padding: 20px 30px; text-align: center; }}
        .stat-card .num {{ font-size: 2.5em; color: #a855f7; font-weight: bold; }}
        .stat-card .label {{ color: #c4b5fd; margin-top: 5px; }}
        .section {{ padding: 30px; max-width: 1000px; margin: 0 auto; }}
        .section h2 {{ color: #a855f7; border-bottom: 1px solid #7c3aed; padding-bottom: 10px; margin-bottom: 20px; }}
        .vuln-card {{ background: #1a0a2e; border-left: 4px solid #ef4444; border-radius: 8px; padding: 20px; margin-bottom: 15px; }}
        .vuln-header {{ display: flex; align-items: center; gap: 15px; margin-bottom: 10px; }}
        .badge {{ padding: 3px 10px; border-radius: 4px; font-size: 0.8em; font-weight: bold; }}
        .badge-danger {{ background: #ef4444; color: white; }}
        .vuln-body p {{ color: #d1d5db; line-height: 1.6; }}
        .vuln-body small {{ color: #6b7280; }}
        .ai-box {{ background: #0f172a; border: 1px solid #3b82f6; border-radius: 10px; padding: 25px; margin-top: 20px; }}
        .ai-box h3 {{ color: #60a5fa; margin-bottom: 15px; }}
        .ai-box p {{ color: #d1d5db; line-height: 1.6; margin-bottom: 10px; }}
        .no-vulns {{ background: #052e16; border: 1px solid #16a34a; border-radius: 8px; padding: 20px; color: #86efac; text-align: center; }}
        .footer {{ text-align: center; padding: 30px; color: #6b7280; border-top: 1px solid #1f2937; margin-top: 40px; }}
        .footer span {{ color: #a855f7; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🛡️ ValkyrieX Pro</h1>
        <div class="author">👩‍💻 Author: <strong>Sonia</strong> 🇵🇰</div>
        <div class="target">🎯 Target: {target}</div>
        <div class="target">🕐 Scan Time: {scan_results.get('scan_time', 'N/A')} UTC</div>
    </div>
    
    <div class="stats">
        <div class="stat-card">
            <div class="num">{len(scan_results.get('module_outputs', []))}</div>
            <div class="label">Modules Run</div>
        </div>
        <div class="stat-card">
            <div class="num" style="color:#ef4444">{len(vulns)}</div>
            <div class="label">Vulnerabilities</div>
        </div>
        <div class="stat-card">
            <div class="num" style="color:#22c55e">{len(safe)}</div>
            <div class="label">Clean Checks</div>
        </div>
    </div>
    
    <div class="section">
        <h2>🐛 Vulnerability Findings</h2>
        {vuln_cards}
        {ai_summary}
    </div>
    
    <div class="footer">
        <p>Generated by <span>ValkyrieX Pro</span> | Author: <span>Sonia 🇵🇰</span></p>
        <p style="margin-top:5px; font-size:0.8em;">For authorized security testing only</p>
    </div>
</body>
</html>"""
    
    with open(path, 'w', encoding='utf-8') as f:
        f.write(html)


def _write_docx(scan_results, ai_analysis, target, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        title = doc.add_heading('ValkyrieX Pro — Security Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run('Author: Sonia 🇵🇰 | Tool: ValkyrieX Pro')
        run.bold = True
        
        doc.add_paragraph()
        
        doc.add_heading('Scan Information', level=1)
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.cell(0, 0).text = 'Target'
        info_table.cell(0, 1).text = target
        info_table.cell(1, 0).text = 'Scan Time'
        info_table.cell(1, 1).text = scan_results.get('scan_time', 'N/A')
        info_table.cell(2, 0).text = 'Report By'
        info_table.cell(2, 1).text = 'Sonia'
        
        doc.add_paragraph()
        
        vulns = [m for m in scan_results.get("module_outputs", []) 
                 if "Vulnerable" in m.get("status", "") or "Alert" in m.get("status", "")]
        
        doc.add_heading(f'Vulnerabilities Found: {len(vulns)}', level=1)
        
        if not vulns:
            doc.add_paragraph('✅ No vulnerabilities detected.')
        else:
            for i, v in enumerate(vulns, 1):
                doc.add_heading(f'{i}. {v["module"].upper().replace("_", " ")}', level=2)
                p = doc.add_paragraph()
                p.add_run('Status: ').bold = True
                p.add_run(v.get('status', 'N/A'))
                p2 = doc.add_paragraph()
                p2.add_run('Summary: ').bold = True
                p2.add_run(v.get('summary', 'N/A'))
                doc.add_paragraph()
        
        if ai_analysis and "ai_analysis" in ai_analysis:
            ai = ai_analysis.get("ai_analysis", {})
            doc.add_heading('AI Analysis', level=1)
            doc.add_paragraph(ai.get('executive_summary', 'N/A'))
            est = doc.add_paragraph()
            est.add_run('Estimated Bounty: ').bold = True
            est.add_run(ai.get('estimated_bounty_range', 'N/A'))
        
        doc.save(path)
        
    except ImportError:
        with open(path.replace('.docx', '_note.txt'), 'w') as f:
            f.write("Install python-docx: pip install python-docx\n")
            f.write("Then re-run the scan to generate DOCX report.\n")
